#!/usr/bin/env python3
"""
스마트 자료취합기 v1.2 — 사무실에서 받는 거의 모든 형식의 "표"를 찾아 하나로 합친다.
--------------------------------------------------------------------------------
입력: .xlsx .xlsm .xltx .xltm .xls(엑셀 97-2003) .csv .tsv .txt
      .docx .doc(워드 97-2003) .pdf(글자형) .hwpx .hwp(한글 5.0)
출력: .xlsx / .docx / .csv / .pdf / .hwpx  (출력 파일의 확장자로 결정)

사용법:
    python smart_merge_v2.py <입력폴더> <출력파일.xlsx|.docx|.csv|.pdf|.hwpx>

v1.0 → v1.1 달라진 점:
- 구버전 한글(.hwp), 엑셀(.xls), 워드(.doc)를 추가 설치 없이 직접 읽는다 (legacy_formats.py).
  암호/배포용 문서, 한글 3.0 등 읽을 수 없는 경우는 이유를 '처리못한파일'에 적어준다.
- 엑셀은 첫 시트만이 아니라 모든 시트를 읽고, 한 시트에 표가 여러 개 있으면 각각 찾는다.
- 표 위의 제목 줄·빈 줄·페이지마다 반복된 머리글을 자동으로 걸러낸다 (모든 형식 공통).
- CSV 는 cp949(한글 윈도우 기본)·UTF-8·UTF-16 을 자동 판별하고, 탭/세미콜론 구분도 읽는다.
- 워드·한글 표 안에 들어있는 표(중첩 표)도 찾는다.
- 하위 폴더까지 읽는 옵션 추가 (include_subfolders).

v1.1 → v1.2 달라진 점 (성능·정확도):
- 파일을 여러 개 동시에 읽는다 (CPU 수만큼 병렬). 파일이 많을수록 빨라진다.
- 두 줄짜리 머리글(예: 위 줄 "연락처" 아래 줄 "본인/보호자")을 "연락처 본인", "연락처 보호자"로 합쳐 읽는다.
- 엑셀 병합 셀(세로로 합쳐진 학과명 등)을 아래 칸까지 채워서 빈칸으로 남지 않게 한다.
- 머리글이 비어 있는데 데이터는 있는 열을 버리지 않고 "(이름 없는 열 N)"으로 살린다.
- "합계/소계/총계" 줄을 자동으로 뺀다 (drop_totals). 똑같은 줄 제거 옵션(dedupe).
- "1,234" 같은 글자 숫자를 진짜 숫자로 바꿔 엑셀에서 바로 계산할 수 있게 한다.
- 결과 엑셀에 머리글 고정·자동 필터·요약 시트를 넣는다.

※ .hwpx로 출력하려면 smart_merge_v2.py와 같은 폴더에 hwpx_template.hwpx 파일이 필요하다.
"""

import sys
import os
import csv
import difflib
import io
import re
import zipfile
import datetime as _dt
import xml.etree.ElementTree as ET
from collections import OrderedDict

ENGINE_VERSION = "1.2.0"
# 화면판/테스트에서 병렬 읽기를 끌 때 바꾼다 ("auto" / True / False)
DEFAULT_PARALLEL = "auto"
from openpyxl import load_workbook, Workbook
from openpyxl.styles import Font, PatternFill
from openpyxl.utils import get_column_letter

# ---------------------------------------------------------------------------
# 1. 동의어 사전 (v0.1과 동일 + 확장 여지)
# ---------------------------------------------------------------------------
SYNONYM_GROUPS = [
    ["이름", "성명", "성 명", "성함", "name", "Name", "한글이름", "한글성명", "학생명", "직원명", "사원명", "회원명", "신청자", "신청자명", "대상자", "대상자명", "참가자", "참가자명"],
    ["영문이름", "영문성명", "영문명", "english name", "영문 이름"],
    ["번호", "순번", "연번", "no", "no.", "num", "순서", "일련번호", "seq", "#"],
    ["학번", "학생번호", "student id", "student no"],
    ["사번", "직원번호", "사원번호", "employee id", "employee no", "직번"],
    ["회원번호", "회원id", "member id", "id", "아이디"],
    ["부서", "소속", "팀", "부서명", "소속부서", "소속명", "팀명", "department", "dept", "team", "소속기관", "기관명", "학과", "학과명", "전공", "소속학과"],
    ["직급", "직위", "직책", "직명", "position", "title", "직급/직위", "직위(직급)"],
    ["전화번호", "연락처", "휴대폰", "휴대전화", "핸드폰", "핸드폰번호", "휴대폰번호", "휴대전화번호", "전화", "phone", "mobile", "tel", "cell", "hp", "h.p", "연락처(휴대폰)", "휴대번호", "연락처_본인", "본인연락처", "휴대폰 번호", "전화 번호", "h.p.", "hp번호"],
    ["사무실전화", "내선", "내선번호", "직장전화", "office phone", "office tel", "유선전화"],
    ["이메일", "메일", "email", "e-mail", "이메일주소", "메일주소", "전자우편"],
    ["생년월일", "출생일자", "생일", "생년", "birth", "birthday", "date of birth", "dob"],
    ["성별", "gender", "sex"],
    ["국적", "국가", "nationality", "country", "출신국"],
    ["주소", "거주지", "address", "주소지", "현주소", "자택주소"],
    ["등록일자", "등록일", "가입일", "가입일자", "신청일", "신청일자", "접수일", "접수일자", "접수일시"],
    ["입사일", "입사일자", "입사년월일", "임용일", "임용일자", "채용일"],
    ["금액", "가격", "단가", "총액", "총금액", "합계금액", "price", "amount", "금액(원)"],
    ["수량", "개수", "qty", "quantity", "인원", "인원수"],
    ["비고", "메모", "특이사항", "remark", "remarks", "note", "notes", "기타", "참고"],
    ["점수", "score", "총점", "합계점수", "평점평균", "평균평점", "gpa"],
    ["등급", "grade"],
    ["상태", "status", "처리상태", "진행상태", "결과", "result"],
    ["여권번호", "passport no", "passport", "passport number"],
    ["외국인등록번호", "등록번호", "외국인번호", "arc no", "arc"],
    ["체류자격", "비자", "visa", "비자종류", "visa type"],
    ["만료일", "만기일", "유효기간", "expiry", "expiration", "체류만료일"],
    ["날짜", "일자", "date", "일시"],
    ["학년", "grade level", "year"],
    ["담당교수", "교수", "교수명", "교원", "담당교원", "지도교수", "강사", "강사명", "professor", "instructor"],
    ["학적상태", "학적", "재학상태", "학적구분"],
    ["입학일", "입학일자", "입학년월일", "입학연도", "입학년도"],
    ["과목", "과목명", "course", "subject", "강좌명", "교과목", "교과목명", "강의명"],
    ["계좌번호", "계좌", "account no", "account number", "입금계좌"],
    ["은행", "은행명", "bank", "거래은행"],
    ["예금주", "예금주명", "account holder"],
]

def build_synonym_lookup():
    lookup = {}
    for group in SYNONYM_GROUPS:
        canonical = group[0]
        for name in group:
            lookup[normalize(name)] = canonical
    return lookup

def normalize(s):
    """열 이름 비교용 정규화: 소문자, 공백/줄바꿈 제거, 괄호 안 설명·특수기호 제거."""
    t = str(s).strip().lower()
    t = re.sub(r"[\s\u3000]+", "", t)          # 공백·줄바꿈
    t = re.sub(r"[*※☆★]", "", t)                 # 필수 표시 기호
    t = re.sub(r"[\(\[（].*?[\)\]）]", "", t) if len(t) > 3 else t   # (선택), (필수) 같은 꼬리표
    t = t.replace("_", "").replace("-", "").replace(".", "").replace(":", "")
    return t

SYN_LOOKUP = build_synonym_lookup()


def _note(*parts):
    """안내 메시지 출력. 콘솔이 한글을 못 찍는 환경(cp1252 등)이나 화면 없는 exe 에서도 절대 죽지 않는다."""
    msg = " ".join(str(p) for p in parts)
    try:
        if sys.stdout is not None:
            sys.stdout.write(msg + "\n")
            sys.stdout.flush()
    except Exception:
        try:
            sys.stdout.buffer.write((msg + "\n").encode("utf-8", "replace"))
        except Exception:
            pass


# ---------------------------------------------------------------------------
# 2. 포맷별 "표 추출기" -- 전부 (헤더, 데이터행들) 리스트를 반환한다.
#    한 파일에서 표가 여러 개 나올 수 있으므로 항상 리스트의 리스트로 반환.
# ---------------------------------------------------------------------------

# ---- 공통: "격자(grid)" 에서 실제 표를 찾아내는 규칙 --------------------------------
# 엑셀 시트/CSV/구버전 문서에서 읽은 값의 격자에는 제목 줄, 빈 줄, 표 여러 개가 섞여 있다.
#  1) 완전히 빈 줄을 경계로 덩어리(block)를 나눈다 (한 시트에 표가 여러 개 있어도 각각 잡는다)
#  2) 덩어리 안에서 "서로 다른 값이 2개 이상 들어있는 첫 줄"을 헤더로 본다
#     (제목 줄처럼 한 칸만 채워진 줄이나 병합된 제목 줄은 자동으로 건너뛴다)
#  3) 헤더와 똑같은 줄이 아래에 또 나오면(페이지마다 반복된 머리글) 지운다
def _is_blank(v):
    return v is None or (isinstance(v, str) and v.strip() == "")


def _cell_str(v):
    if v is None:
        return ""
    if isinstance(v, float) and v.is_integer():
        return str(int(v))
    return str(v).strip()


_NUMBERISH = re.compile(r"^[\s\-+]?[\d,]+(\.\d+)?\s*%?$")


def _looks_like_data_cell(v):
    """숫자·날짜처럼 '값'으로 보이는 칸이면 True (머리글 줄인지 판단할 때 쓴다)."""
    if v is None:
        return False
    if isinstance(v, (int, float, _dt.date, _dt.datetime, _dt.time)):
        return True
    s = str(v).strip()
    if not s:
        return False
    if _NUMBERISH.match(s):
        return True
    if re.match(r"^\d{2,4}[.\-/]\d{1,2}[.\-/]\d{1,2}", s):
        return True
    if re.match(r"^0\d{1,2}-\d{3,4}-\d{4}$", s):   # 전화번호
        return True
    return False


def _looks_like_strong_data(v):
    """머리글에는 거의 안 나오는 값: 날짜형, 전화번호, 5자리 이상 숫자(학번·사번), 소수. (연도 4자리는 제외)"""
    if v is None:
        return False
    if isinstance(v, (_dt.date, _dt.datetime)):
        return True
    if isinstance(v, float) and not v.is_integer():
        return True
    if isinstance(v, (int, float)):
        return abs(int(v)) >= 10000
    s = str(v).strip()
    if re.match(r"^\d{5,}$", s) or re.match(r"^0\d{1,2}-\d{3,4}-\d{4}$", s):
        return True
    if re.match(r"^\d{2,4}[.\-/]\d{1,2}[.\-/]\d{1,2}", s) or re.match(r"^\d{6}-\d", s):
        return True
    return False


def _combine_two_row_header(top, sub):
    """두 줄 머리글을 한 줄로 합친다. 위 줄이 병합돼 같은 이름이 반복되면 '위 아래' 꼴로 만든다."""
    out = []
    for t, s in zip(top, sub):
        t, s = _cell_str(t), _cell_str(s)
        if t and s and t != s:
            out.append(f"{t} {s}")
        else:
            out.append(t or s)
    return out


def _is_sub_header_row(header, row, min_cols=2):
    """머리글 바로 아래 줄이 '작은 머리글'(2줄 머리글의 아래 줄)인지 판단.
    조건: 값으로 보이는 칸(숫자·날짜·전화번호)이 하나도 없고, 글자 칸이 2개 이상이며,
          위 머리글에 같은 이름이 반복되거나(가로 병합) 빈 칸이 있을 것."""
    if row is None:
        return False
    cells = [_cell_str(v) for v in row]
    filled = [c for c in cells if c]
    if len(filled) < min_cols:
        return False
    if any(_looks_like_data_cell(v) for v in row):
        return False
    hdr = [_cell_str(h) for h in header]
    non_blank = [h for h in hdr if h]
    has_dup = len(non_blank) != len(set(non_blank))
    has_gap = any(not h for h in hdr)
    if not (has_dup or has_gap):
        return False
    # 아래 줄이 위 머리글과 완전히 같으면 그냥 반복 머리글이다.
    # (세로 병합으로 '학과/이름/비고'가 아래 줄에도 복사된 경우는 나머지 칸이 다르므로 통과)
    same = sum(1 for h, c in zip(hdr, cells) if h and c and h == c)
    if same >= len(filled):
        return False
    # 위 머리글이 중복/빈칸인 열 중 최소 하나는 아래 줄에 새 글자가 있어야 한다
    dup_names = {h for h in non_blank if non_blank.count(h) > 1}
    fresh = any((not h or h in dup_names) and c and c != h for h, c in zip(hdr, cells))
    return fresh


def split_grid_into_tables(grid, min_cols=2):
    tables = []
    block = []

    def flush(block):
        # 헤더 찾기
        header_idx = None
        for i, row in enumerate(block):
            vals = [_cell_str(v) for v in row if not _is_blank(v)]
            if len(vals) >= min_cols and len(set(vals)) >= min_cols:
                header_idx = i
                break
        if header_idx is None:
            return
        header_raw = list(block[header_idx])
        # 머리글 왼쪽/오른쪽의 완전히 빈 열은 잘라낸다. 단, 머리글은 비었지만 아래 데이터가
        # 들어있는 열은 살린다 ("(이름 없는 열 N)" 으로 표시됨).
        body = block[header_idx + 1:]
        width = max([len(header_raw)] + [len(r) for r in body])
        header_raw += [None] * (width - len(header_raw))
        used = [not _is_blank(v) for v in header_raw]
        for r in body:
            for i, v in enumerate(r):
                if not _is_blank(v):
                    used[i] = True
        if not any(used):
            return
        first = used.index(True)
        last = len(used) - 1 - used[::-1].index(True)
        header = [_cell_str(v) for v in header_raw[first:last + 1]]
        # 머리글 줄이 없는 표(첫 줄부터 바로 사람 데이터인 경우): 머리글 후보에 전화번호·날짜·
        # 긴 번호(5자리 이상) 같은 '값'이 2개 이상이면 머리글이 아니라고 보고, 첫 줄도 데이터로 살린다.
        strong = sum(1 for v in header_raw[first:last + 1] if _looks_like_strong_data(v))
        if strong >= 2 and body:
            body = [header_raw] + body
            header = [""] * (last - first + 1)
        # 두 줄 머리글이면 합친다
        elif body and _is_sub_header_row(header, list(body[0])[first:last + 1] + [None] * 0, min_cols):
            sub = list(body[0]) + [None] * width
            header = _combine_two_row_header(header, sub[first:last + 1])
            body = body[1:]
        data = []
        for row in body:
            cells = list(row[first:last + 1]) + [None] * max(0, (last - first + 1) - len(row[first:last + 1]))
            if all(_is_blank(v) for v in cells):
                continue
            if [_cell_str(v) for v in cells] == header:   # 반복된 머리글
                continue
            data.append(cells)
        if data:
            tables.append((header, data))

    for row in grid:
        if row is None or all(_is_blank(v) for v in row):
            if block:
                flush(block)
                block = []
        else:
            block.append(list(row))
    if block:
        flush(block)
    return tables


def _tidy_table(header, data):
    """문서(워드/한글/PDF)에서 나온 표 하나에도 같은 규칙을 적용한다: 제목 줄·빈 줄·반복 머리글 제거."""
    grid = [list(header)] + [list(r) for r in data]
    found = split_grid_into_tables(grid)
    return found


_MERGE_RE = re.compile(r'<mergeCell\s+ref="([A-Z]+)(\d+):([A-Z]+)(\d+)"')


def _col_index(letters):
    n = 0
    for ch in letters:
        n = n * 26 + (ord(ch) - 64)
    return n - 1


def _xlsx_merged_ranges(wb, ws):
    """read_only 모드에는 병합 정보가 없어서 시트 XML에서 직접 읽는다. [(r1,c1,r2,c2) 0기준]"""
    try:
        xml = wb._archive.read(ws._worksheet_path).decode("utf-8", "ignore")
    except Exception:
        return []
    out = []
    for c1, r1, c2, r2 in _MERGE_RE.findall(xml):
        out.append((int(r1) - 1, _col_index(c1), int(r2) - 1, _col_index(c2)))
    return out


def _fill_merged(rows, ranges):
    """병합 범위의 왼쪽 위 값을 범위 전체 칸에 복사한다 (세로 병합된 학과명, 가로 병합된 머리글)."""
    for r1, c1, r2, c2 in ranges:
        if r1 >= len(rows):
            continue
        row0 = rows[r1]
        if c1 >= len(row0):
            continue
        val = row0[c1]
        if _is_blank(val):
            continue
        for r in range(r1, min(r2, len(rows) - 1) + 1):
            row = rows[r]
            if len(row) <= c2:
                row.extend([None] * (c2 + 1 - len(row)))
            for c in range(c1, c2 + 1):
                if _is_blank(row[c]):
                    row[c] = val
    return rows


def extract_tables_xlsx(path):
    wb = load_workbook(path, data_only=True, read_only=True)
    tables = []
    try:
        for ws in wb.worksheets:
            if ws.sheet_state != "visible":
                continue
            rows = [list(row) for row in ws.iter_rows(min_row=1, values_only=True)]
            ranges = _xlsx_merged_ranges(wb, ws)
            if ranges:
                rows = _fill_merged(rows, ranges)
            tables.extend(split_grid_into_tables(rows))
    finally:
        wb.close()
    return tables


def _read_text_any_encoding(path):
    with open(path, "rb") as f:
        raw = f.read()
    if raw.startswith(b"\xff\xfe") or raw.startswith(b"\xfe\xff"):
        return raw.decode("utf-16", errors="replace")
    for enc in ("utf-8-sig", "cp949", "euc-kr", "utf-16"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def extract_tables_csv(path):
    text = _read_text_any_encoding(path)
    sample = text[:4096]
    try:
        dialect = csv.Sniffer().sniff(sample, delimiters=",\t;|")
    except csv.Error:
        dialect = csv.excel
        if path.lower().endswith((".tsv", ".txt")) and "\t" in sample:
            dialect = csv.excel_tab
    rows = list(csv.reader(io.StringIO(text), dialect))
    return split_grid_into_tables(rows)


def extract_tables_docx(path):
    import docx
    doc = docx.Document(path)
    tables = []

    def walk(tbl_list):
        for tbl in tbl_list:
            grid = []
            for row in tbl.rows:
                cells = []
                prev = None
                for cell in row.cells:
                    # 가로로 병합된 셀은 python-docx 가 같은 셀 객체를 반복해서 준다
                    cells.append(cell.text.strip())
                    prev = cell
                grid.append(cells)
                for cell in row.cells:
                    if cell.tables:
                        walk(cell.tables)
            if grid:
                tables.extend(split_grid_into_tables(grid))

    walk(doc.tables)
    return tables
def _extract_table_by_position(page, x_gap=15, y_tol=3):
    """
    테두리(그리드선)가 없는 PDF 표를 위한 대체 추출기.
    글자들의 x좌표(가로 위치)를 기준으로 열을 추정한다.
    (워드 -> PDF 변환처럼 표에 테두리가 없는 경우 pdfplumber의 기본 표 인식이 실패하기 때문)
    """
    words = page.extract_words()
    if not words:
        return None
    # 같은 줄(y좌표가 비슷한) 단어들끼리 묶기
    lines = []
    for w in sorted(words, key=lambda w: (w["top"], w["x0"])):
        placed = False
        for line in lines:
            if abs(line[0]["top"] - w["top"]) <= y_tol:
                line.append(w)
                placed = True
                break
        if not placed:
            lines.append([w])
    lines = [sorted(l, key=lambda w: w["x0"]) for l in lines]
    if len(lines) < 2:
        return None

    def build_columns(header_words):
        columns = []
        for w in header_words:
            if columns and (w["x0"] - columns[-1]["x1"]) < x_gap:
                columns[-1]["label"] += w["text"]
                columns[-1]["x1"] = w["x1"]
            else:
                columns.append({"x0": w["x0"], "x1": w["x1"], "label": w["text"]})
        return columns

    # 제목/설명 줄이 표 위에 섞여 있을 수 있으므로, 열이 2개 이상 잡히는 첫 줄을
    # 실제 헤더로 간주한다 (제목처럼 한 덩어리로 뭉치는 줄은 건너뜀).
    header_idx = None
    columns = None
    for i, line in enumerate(lines[:-1]):
        cols = build_columns(line)
        if len(cols) >= 2:
            header_idx, columns = i, cols
            break
    if columns is None:
        return None

    col_starts = [c["x0"] for c in columns]
    header = [c["label"] for c in columns]

    data_rows = []
    for line in lines[header_idx + 1:]:
        row = [""] * len(columns)
        for w in line:
            idx = min(range(len(col_starts)), key=lambda i: abs(col_starts[i] - w["x0"]))
            row[idx] = (row[idx] + " " + w["text"]).strip() if row[idx] else w["text"]
        if any(v for v in row):
            data_rows.append(row)
    if not data_rows:
        return None
    return (header, data_rows)


_NOTE_PREFIX = ("※", "*", "＊", "주)", "주:", "출처")


def _join_wrapped(a, b):
    """줄바꿈으로 끊긴 글자를 다시 붙인다. 한글끼리는 붙여 쓰고(글로벌호스피탈리+티산업), 그 외엔 띄어 쓴다."""
    a, b = str(a).rstrip(), str(b).lstrip()
    if not a:
        return b
    if not b:
        return a
    if re.match(r"[가-힣]", a[-1]) and re.match(r"[가-힣(]", b[0]) and not a.endswith((",", ".", ")")):
        return a + b
    return a + " " + b


def merge_wrapped_rows(header, rows):
    """PDF 처럼 '칸 안에서 줄바꿈된 글'이 다음 줄로 떨어져 나온 표를 원래 줄로 되돌린다.
    조건: 첫 칸이 비어 있고, 글자가 있는 칸이 모두 '바로 윗줄에도 글자가 있던 칸'이며, 값(숫자·날짜 등)이 아닐 것.
    '※ …' 처럼 표 아래 붙은 각주 줄은 뺀다."""
    out = []
    after_note = False
    for row in rows:
        cells = list(row)
        first = next((c for c in cells if not _is_blank(c)), None)
        if first is not None and str(first).strip().startswith(_NOTE_PREFIX):
            after_note = True
            continue
        if after_note:
            # 각주가 여러 줄로 이어진 경우: 값(숫자 등)도 없고 칸도 절반 이상 비어 있으면 각주의 연속으로 본다
            filled_n = sum(1 for c in cells if not _is_blank(c))
            if filled_n * 2 <= len(cells) and not any(_looks_like_strong_data(c) for c in cells):
                continue
            after_note = False
        if out and _is_blank(cells[0]):
            prev = out[-1]
            filled = [i for i, c in enumerate(cells) if not _is_blank(c)]
            if filled and all(i < len(prev) and not _is_blank(prev[i]) for i in filled) \
                    and not any(_looks_like_strong_data(cells[i]) for i in filled):
                for i in filled:
                    prev[i] = _join_wrapped(prev[i], cells[i])
                continue
        out.append(cells)
    return out


def extract_tables_pdf(path):
    import pdfplumber
    tables = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            found = page.extract_tables()
            page_tables = []
            for tbl in found:
                if not tbl:
                    continue
                page_tables.extend(split_grid_into_tables([[("" if c is None else str(c).strip()) for c in r] for r in tbl]))
            if not page_tables:
                # 테두리 없는 표(예: 워드->PDF 변환본) 대비 위치기반 추출 시도
                fallback = _extract_table_by_position(page)
                if fallback:
                    page_tables.extend(_tidy_table(*fallback))
            tables.extend((h, merge_wrapped_rows(h, r)) for h, r in page_tables)
    return [(h, r) for h, r in tables if r]

def _tag(el):
    """네임스페이스 접두어를 뗀 태그 이름 ('{ns}tbl' -> 'tbl')."""
    return el.tag.rsplit("}", 1)[-1]


def _cell_text(tc):
    """
    <hp:tc> 안의 모든 <hp:t> 텍스트를 이어붙인다 (문단이 여러 개면 줄바꿈으로 구분).
    빈 문단(세로쓰기 글자 사이의 간격용 빈 줄 등)은 제거해 불필요한 줄바꿈 잡음을 없앤다.
    """
    parts = []
    for p in tc.iter():
        if _tag(p) == "p":
            line = "".join(t.text or "" for t in p.iter() if _tag(t) == "t").strip()
            if line:
                parts.append(line)
    return "\n".join(parts).strip()


def _parse_hwpx_table(tbl_el):
    """
    <hp:tbl> 하나를 colAddr/rowAddr/colSpan/rowSpan 정보를 이용해
    실제 그리드(빈칸 없는 표)로 복원한다. (병합된 셀 때문에 열이 밀리는 문제 방지)
    병합된 영역은 왼쪽위 셀의 텍스트로 채운다(중복되더라도 열 정렬을 지키는 것이 우선).
    """
    row_cnt = int(tbl_el.attrib.get("rowCnt", 0) or 0)
    col_cnt = int(tbl_el.attrib.get("colCnt", 0) or 0)

    cells = []
    for tr in tbl_el:
        if _tag(tr) != "tr":
            continue
        for tc in tr:
            if _tag(tc) != "tc":
                continue
            addr = next((c for c in tc if _tag(c) == "cellAddr"), None)
            span = next((c for c in tc if _tag(c) == "cellSpan"), None)
            if addr is None:
                continue
            col0 = int(addr.attrib.get("colAddr", 0))
            row0 = int(addr.attrib.get("rowAddr", 0))
            col_span = int(span.attrib.get("colSpan", 1)) if span is not None else 1
            row_span = int(span.attrib.get("rowSpan", 1)) if span is not None else 1
            cells.append((row0, col0, row_span, col_span, _cell_text(tc)))

    if not cells:
        return None
    if row_cnt == 0:
        row_cnt = max(r + rs for r, c, rs, cs, t in cells)
    if col_cnt == 0:
        col_cnt = max(c + cs for r, c, rs, cs, t in cells)

    grid = [["" for _ in range(col_cnt)] for _ in range(row_cnt)]
    for row0, col0, row_span, col_span, text in cells:
        for r in range(row0, min(row0 + row_span, row_cnt)):
            for c in range(col0, min(col0 + col_span, col_cnt)):
                grid[r][c] = text

    if row_cnt < 2:
        return None
    return grid[0], grid[1:]


def extract_tables_hwpx(path):
    """
    HWPX(zip+XML) 표 추출. 실제 .hwpx 파일(신경주대학교 한국어학당 신입생 서류
    심사표.hwpx)로 검증 완료 -- colAddr/rowAddr/cellSpan 기반으로 병합 셀이 있어도
    열이 밀리지 않도록 그리드를 복원한다.
    """
    tables = []
    with zipfile.ZipFile(path) as z:
        section_files = [n for n in z.namelist() if n.startswith("Contents/section") and n.endswith(".xml")]
        for sec in sorted(section_files):
            xml_bytes = z.read(sec)
            root = ET.fromstring(xml_bytes)
            for tbl_el in root.iter():
                if _tag(tbl_el) == "tbl":
                    parsed = _parse_hwpx_table(tbl_el)
                    if parsed:
                        tables.extend(_tidy_table(*parsed))
    return tables


def extract_tables_xls(path):
    from legacy_formats import extract_grids_xls
    tables = []
    for name, grid in extract_grids_xls(path):
        tables.extend(split_grid_into_tables(grid))
    return tables


def extract_tables_hwp(path):
    from legacy_formats import extract_tables_hwp as _hwp
    tables = []
    for header, data in _hwp(path):
        tables.extend(_tidy_table(header, data))
    return tables


def extract_tables_doc(path):
    from legacy_formats import extract_tables_doc as _doc
    tables = []
    for header, data in _doc(path):
        tables.extend(_tidy_table(header, data))
    return tables


EXTRACTORS = {
    ".xlsx": extract_tables_xlsx,
    ".xlsm": extract_tables_xlsx,
    ".xltx": extract_tables_xlsx,
    ".xltm": extract_tables_xlsx,
    ".xls": extract_tables_xls,
    ".csv": extract_tables_csv,
    ".tsv": extract_tables_csv,
    ".txt": extract_tables_csv,
    ".docx": extract_tables_docx,
    ".doc": extract_tables_doc,
    ".pdf": extract_tables_pdf,
    ".hwpx": extract_tables_hwpx,
    ".hwp": extract_tables_hwp,
}

# 형식은 알지만 아직 읽지 못하는 것들에 대한 안내문
UNSUPPORTED_HINT = {
    ".pptx": "파워포인트 파일은 지원하지 않습니다. 표를 엑셀이나 워드로 옮겨 저장한 뒤 넣어주세요.",
    ".ppt": "파워포인트 파일은 지원하지 않습니다. 표를 엑셀이나 워드로 옮겨 저장한 뒤 넣어주세요.",
    ".odt": "오픈오피스 문서(.odt)는 지원하지 않습니다. 워드(.docx)로 저장한 뒤 넣어주세요.",
    ".ods": "오픈오피스 표(.ods)는 지원하지 않습니다. 엑셀(.xlsx)로 저장한 뒤 넣어주세요.",
    ".jpg": "그림 파일은 표를 읽을 수 없습니다.",
    ".jpeg": "그림 파일은 표를 읽을 수 없습니다.",
    ".png": "그림 파일은 표를 읽을 수 없습니다.",
}


# ---------------------------------------------------------------------------
# 3. 열 매핑 (v0.1과 동일)
# ---------------------------------------------------------------------------
def map_column(col_name, canonical_columns, similarity_threshold=0.85):
    """열 이름 하나를 표준 열 이름으로 바꾼다.
    1) 동의어 사전에 있으면 사전매칭
    2) "부서명"→"부서", "항목명"→"항목" 처럼 '명/값/란' 꼬리만 다른 경우
    3) 이미 나온 표준 열과 글자가 거의 같은 경우(0.85 이상)만 유사도매칭 —
       '번호'≠'전화번호', '평가등급'≠'평가점수' 처럼 뜻이 다른 열이 합쳐지는 사고를 막기 위해 엄격하게 본다.
    4) 그 외에는 새 열로 두고 '수동확인필요'로 표시한다."""
    norm = normalize(col_name)
    if norm in SYN_LOOKUP:
        return SYN_LOOKUP[norm], "사전매칭", 1.0
    for suffix in ("명", "값", "란", "칸"):
        if len(norm) > 2 and norm.endswith(suffix):
            base = norm[:-1]
            if base in SYN_LOOKUP:
                return SYN_LOOKUP[base], "사전매칭", 1.0
            for canon in canonical_columns:
                if normalize(canon) == base:
                    return canon, "유사도매칭", 0.95
    best_match, best_score = None, 0.0
    for canon in canonical_columns:
        cn = normalize(canon)
        if cn == norm:
            return canon, "사전매칭", 1.0
        if min(len(cn), len(norm)) < 2:
            continue
        score = difflib.SequenceMatcher(None, norm, cn).ratio()
        # 한쪽이 다른 쪽을 통째로 포함하고 길이 차이가 1글자뿐이면(예: 연락처/연락처1) 같은 열로 본다
        if (norm in cn or cn in norm) and abs(len(cn) - len(norm)) <= 1 and min(len(cn), len(norm)) >= 2:
            score = max(score, 0.9)
        if score > best_score:
            best_match, best_score = canon, score
    if best_score >= similarity_threshold:
        return best_match, "유사도매칭", round(best_score, 2)
    return str(col_name).strip(), "신규열(수동확인필요)", 0.0


# ---------------------------------------------------------------------------
# 4. 메인 병합 로직 -- 입력 폴더를 읽어 (표준열목록, 합쳐진행, 매핑정보, 못읽은파일) 반환.
#    출력 형식(엑셀/워드/CSV)과는 분리되어 있다 -- 아래 5번 참고.
# ---------------------------------------------------------------------------
def iter_input_files(input_dir, include_subfolders=False):
    """입력 폴더 안의 파일 경로를 (표시이름, 경로) 로 낸다. 임시파일(~$, .)은 제외."""
    if include_subfolders:
        for root, dirs, files in os.walk(input_dir):
            dirs.sort()
            for fname in sorted(files):
                if fname.startswith("~$") or fname.startswith("."):
                    continue
                rel = os.path.relpath(os.path.join(root, fname), input_dir)
                yield rel, os.path.join(root, fname)
    else:
        for fname in sorted(os.listdir(input_dir)):
            if fname.startswith("~$") or fname.startswith("."):
                continue
            path = os.path.join(input_dir, fname)
            if os.path.isfile(path):
                yield fname, path


_TOTAL_WORDS = {"합계", "총계", "소계", "계", "총합", "합", "총합계", "total", "sum", "subtotal"}


def _is_total_row(row):
    """첫 번째로 값이 들어있는 칸이 '합계/소계/총계/계' 이면 합계 줄로 본다."""
    for v in row:
        if _is_blank(v):
            continue
        s = str(v).strip().replace(" ", "").lower()
        return s in _TOTAL_WORDS
    return False


_INT_RE = re.compile(r"^-?\d{1,15}$")
_NUM_COMMA_RE = re.compile(r"^-?\d{1,3}(,\d{3})+(\.\d+)?$")
_FLOAT_RE = re.compile(r"^-?\d+\.\d+$")


def coerce_value(v):
    """'1,234' → 1234, '3.5' → 3.5 처럼 글자로 된 숫자를 진짜 숫자로. '0123'(앞자리 0)·전화번호·날짜는 그대로 둔다."""
    if not isinstance(v, str):
        return v
    s = v.strip()
    if not s or len(s) > 20:
        return v
    if _INT_RE.match(s):
        if len(s.lstrip("-")) > 1 and s.lstrip("-")[0] == "0":
            return v    # 앞자리 0 은 번호(학번·코드)일 가능성이 높다
        try:
            return int(s)
        except ValueError:
            return v
    if _NUM_COMMA_RE.match(s):
        try:
            n = float(s.replace(",", ""))
            return int(n) if n.is_integer() and "." not in s else n
        except ValueError:
            return v
    if _FLOAT_RE.match(s):
        try:
            return float(s)
        except ValueError:
            return v
    return v


def _extract_one(job):
    """작업자 프로세스에서 파일 하나를 읽는다. 결과는 pickle 가능한 값만 돌려준다."""
    fname, path, ext = job
    try:
        tables = EXTRACTORS[ext](path)
        return fname, tables, None
    except Exception as e:
        from legacy_formats import LegacyFormatError
        if isinstance(e, LegacyFormatError):
            return fname, None, str(e)
        return fname, None, f"읽기 실패: {type(e).__name__}: {e}"


def _run_extractions(jobs, progress=None, parallel="auto"):
    """파일들을 읽어 [(fname, tables, error)] 를 입력 순서대로 돌려준다.
    parallel: "auto"(파일 4개 이상이고 CPU 2개 이상이면 병렬) / True / False"""
    n = len(jobs)
    use_parallel = False
    if parallel is True or parallel == "auto":
        try:
            import multiprocessing
            cpus = multiprocessing.cpu_count()
        except Exception:
            cpus = 1
        if parallel is True or (n >= 4 and cpus >= 2):
            use_parallel = sys.platform != "emscripten" and cpus >= 2
    results = [None] * n
    if use_parallel:
        try:
            from concurrent.futures import ProcessPoolExecutor, as_completed
            workers = max(2, min(8, cpus - 1 if cpus > 2 else cpus))
            done = 0
            with ProcessPoolExecutor(max_workers=workers) as ex:
                fut_map = {ex.submit(_extract_one, job): i for i, job in enumerate(jobs)}
                for fut in as_completed(fut_map):
                    i = fut_map[fut]
                    results[i] = fut.result()   # 작업자 풀 자체가 깨지면 예외 → 아래에서 한 개씩 다시 읽는다
                    done += 1
                    if progress:
                        progress(f"읽는 중… {done}/{n}  ({jobs[i][0]})")
            return results
        except Exception:
            results = [None] * n   # 병렬이 안 되는 환경이면 조용히 한 개씩 읽는다
            if progress:
                progress("병렬 읽기가 안 되는 환경이라 한 개씩 읽어요…")
    for i, job in enumerate(jobs):
        if progress:
            progress(f"읽는 중… {i + 1}/{n}  ({job[0]})")
        results[i] = _extract_one(job)
    return results


def merge_all(input_dir, include_subfolders=False, progress=None,
              drop_totals=True, dedupe=False, parallel=None, stats=None):
    """progress(메시지) 콜백을 주면 파일마다 진행 상황을 알려준다.
    drop_totals: '합계/소계/총계' 줄을 뺀다.   dedupe: 완전히 똑같은 줄은 한 번만 남긴다.
    stats: dict 를 주면 파일 수·표 수·뺀 줄 수 등을 채워 준다."""
    canonical_columns = []
    file_mappings = []   # (표시이름, mapping dict)
    all_rows = []         # (row_dict, 표시이름)
    skipped = []
    if parallel is None:
        parallel = DEFAULT_PARALLEL
    st = stats if stats is not None else {}
    st.update({"files": 0, "tables": 0, "rows": 0, "totals_dropped": 0, "dupes_dropped": 0, "unnamed_cols": 0})

    jobs = []
    for fname, path in iter_input_files(input_dir, include_subfolders):
        ext = os.path.splitext(fname)[1].lower()
        if ext in UNSUPPORTED_HINT:
            skipped.append((fname, UNSUPPORTED_HINT[ext]))
            continue
        if ext not in EXTRACTORS:
            continue
        jobs.append((fname, path, ext))
    st["files"] = len(jobs)

    seen_rows = set()
    for fname, tables, err in _run_extractions(jobs, progress, parallel):
        if err is not None:
            skipped.append((fname, err))
            continue
        if not tables:
            skipped.append((fname, "표를 찾지 못함"))
            continue

        for idx, (header, data) in enumerate(tables):
            label = fname if len(tables) == 1 else f"{fname} (표{idx+1})"
            st["tables"] += 1

            # 같은 표 안에 이름이 똑같은 열이 여러 개 있으면(예: 평가등급 A/B/C/D/E가
            # 전부 "평가등급"으로만 적힌 병합표) 열 이름으로만 매칭할 경우 데이터가
            # 서로 덮어써서 사라진다. 위치(열 순서) 기준으로 구분해 각각 별도 열로 취급한다.
            # 머리글이 비어 있는데 데이터가 있는 열은 "(이름 없는 열 N)" 으로 살린다.
            seen_count = {}
            clean_header = []
            unnamed = []
            for ci, col in enumerate(header):
                col = "" if col is None else str(col).strip()
                if col == "":
                    if any(ci < len(r) and not _is_blank(r[ci]) for r in data):
                        col = f"(이름 없는 열 {ci + 1})"
                        unnamed.append(col)
                    else:
                        clean_header.append("")
                        continue
                seen_count[col] = seen_count.get(col, 0) + 1
                clean_header.append(col if seen_count[col] == 1 else f"{col} ({seen_count[col]})")
            st["unnamed_cols"] += len(unnamed)

            # 위치별로 표준 열을 매핑 (열 이름이 아니라 순서 기준이라 중복 이름도 안전).
            # 단, 같은 표 안에서 이름이 중복된 열(예: "평가등급 (2)")은 유사도 매칭을
            # 시키면 원래 열("평가등급")과 다시 합쳐져 버리므로, 그런 경우는 무조건
            # 새 열로 등록해서 데이터가 서로 덮어쓰지 않게 한다.
            position_mapping = []
            mapping_for_report = OrderedDict()
            for original, col in zip(header, clean_header):
                if col == "":
                    position_mapping.append(None)
                    continue
                orig_s = "" if original is None else str(original).strip()
                if col in unnamed:
                    canon, method, score = col, "신규열(수동확인필요)", 0.0
                elif col != orig_s:
                    canon, method, score = col, "표내중복열(수동확인필요)", 0.0
                else:
                    canon, method, score = map_column(col, canonical_columns)
                if canon not in canonical_columns:
                    canonical_columns.append(canon)
                position_mapping.append(canon)
                mapping_for_report[col] = (canon, method, score)
            file_mappings.append((label, mapping_for_report))

            for row in data:
                if drop_totals and _is_total_row(row):
                    st["totals_dropped"] += 1
                    continue
                row_dict = {}
                for i, val in enumerate(row):
                    if i >= len(position_mapping) or position_mapping[i] is None:
                        continue
                    row_dict[position_mapping[i]] = coerce_value(val)
                if not any(v not in (None, "") for v in row_dict.values()):
                    continue
                if dedupe:
                    key = tuple((k, str(v)) for k, v in sorted(row_dict.items()) if v not in (None, ""))
                    if key in seen_rows:
                        st["dupes_dropped"] += 1
                        continue
                    seen_rows.add(key)
                all_rows.append((row_dict, label))

    st["rows"] = len(all_rows)
    return canonical_columns, all_rows, file_mappings, skipped


# ---------------------------------------------------------------------------
# 5. 출력 형식별 저장기 -- 확장자만 다르게 주면 원하는 파일로 결과를 받을 수 있다.
# ---------------------------------------------------------------------------
# 화면/결과 파일에 보여줄 때 쓰는 쉬운 말 (내부 판단 문자열은 그대로 두고 표시만 바꾼다)
METHOD_LABEL = {
    "사전매칭": "같은 뜻 (자동)",
    "유사도매칭": "비슷한 이름 (자동)",
    "신규열(수동확인필요)": "새 열 — 확인 필요",
    "표내중복열(수동확인필요)": "같은 이름이 여러 개 — 확인 필요",
}
MAPPING_SHEET = "열이름맞춤표"
MAPPING_HEADER = ["파일(표)", "원래 열 이름", "→ 합쳐진 열 이름", "어떻게 맞췄나", "확신도"]


def _method_label(m):
    return METHOD_LABEL.get(m, m)


def write_xlsx(output_path, canonical_columns, all_rows, file_mappings, skipped):
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "통합결과"
    header_row = canonical_columns + ["출처파일"]
    ws1.append(header_row)
    for cell in ws1[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="DCE6F1")
    for row_dict, source in all_rows:
        ws1.append([row_dict.get(c, "") for c in canonical_columns] + [source])
    ws1.freeze_panes = "A2"
    if all_rows:
        ws1.auto_filter.ref = ws1.dimensions

    ws2 = wb.create_sheet(MAPPING_SHEET)
    ws2.append(MAPPING_HEADER)
    for cell in ws2[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="DCE6F1")
    for label, mapping in file_mappings:
        for orig, (canon, method, score) in mapping.items():
            ws2.append([label, orig, canon, _method_label(method), score])
            if method != "사전매칭":
                for c in ws2[ws2.max_row]:
                    c.fill = PatternFill("solid", fgColor="FFF2CC")

    if skipped:
        ws3 = wb.create_sheet("처리못한파일")
        ws3.append(["파일명", "읽지 못한 이유 / 안내"])
        for cell in ws3[1]:
            cell.font = Font(bold=True)
            cell.fill = PatternFill("solid", fgColor="F4CCCC")
        for fname, reason in skipped:
            ws3.append([fname, reason])

    # 요약 시트: 어떤 파일에서 몇 줄이 왔는지 한눈에
    ws4 = wb.create_sheet("요약")
    ws4.append(["출처파일(표)", "줄 수"])
    for cell in ws4[1]:
        cell.font = Font(bold=True)
        cell.fill = PatternFill("solid", fgColor="DCE6F1")
    counts = OrderedDict()
    for _, source in all_rows:
        counts[source] = counts.get(source, 0) + 1
    for source, n in counts.items():
        ws4.append([source, n])
    ws4.append(["합계", len(all_rows)])
    ws4[ws4.max_row][0].font = Font(bold=True)
    ws4[ws4.max_row][1].font = Font(bold=True)
    ws4.append([])
    ws4.append(["합쳐진 표 수", len(file_mappings)])
    ws4.append(["열 수", len(canonical_columns)])
    ws4.append(["읽지 못한 파일 수", len(skipped)])
    ws4.append(["만든 프로그램", f"스마트 자료취합기 v{ENGINE_VERSION}"])

    for ws in wb.worksheets:
        widths = {}
        for r_idx, row in enumerate(ws.iter_rows(values_only=True)):
            if r_idx > 500:   # 폭 계산은 앞 500줄만 (큰 결과도 빠르게 저장)
                break
            for c_idx, v in enumerate(row):
                if v is None:
                    continue
                n = len(str(v))
                if n > widths.get(c_idx, 0):
                    widths[c_idx] = n
        for c_idx, n in widths.items():
            ws.column_dimensions[get_column_letter(c_idx + 1)].width = min(max(n + 2, 10), 40)

    wb.save(output_path)


def write_csv(output_path, canonical_columns, all_rows, file_mappings, skipped):
    with open(output_path, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.writer(f)
        writer.writerow(canonical_columns + ["출처파일"])
        for row_dict, source in all_rows:
            writer.writerow([row_dict.get(c, "") for c in canonical_columns] + [source])
    # CSV는 시트를 여러 개 가질 수 없어서 매핑리포트/처리못한파일은 별도 안내로 대체한다.
    if file_mappings or skipped:
        _note("참고: CSV는 표 1개만 담을 수 있어서 열이름맞춤표·처리못한파일 내역은 CSV에는")
        _note("      들어가지 않았습니다. 그 내역까지 확인하려면 .xlsx로 출력해주세요.")


def write_docx(output_path, canonical_columns, all_rows, file_mappings, skipped):
    import docx
    from docx.shared import Pt, Twips
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    def set_table_fixed_layout(table):
        # 워드 기본값은 표를 페이지 너비에 맞춰 열을 "똑같은 너비"로 억지로 나누기 때문에
        # 열이 많으면 한글 한 글자씩 세로로 쪼개져 보이는 문제가 생긴다.
        # 열마다 필요한 너비를 직접 지정하는 "고정 레이아웃"으로 바꿔서 이를 방지한다.
        tbl_pr = table._tbl.tblPr
        layout = tbl_pr.makeelement(qn("w:tblLayout"), {qn("w:type"): "fixed"})
        tbl_pr.append(layout)
        table.autofit = False

    def set_col_widths(table, widths_twips):
        for row in table.rows:
            for cell, w in zip(row.cells, widths_twips):
                cell.width = Twips(w)

    def compute_col_widths(headers, rows, sample=100, min_w=650, max_w=2400, per_char=130, pad=120):
        widths = []
        for i, h in enumerate(headers):
            longest = len(str(h))
            for r in rows[:sample]:
                if i < len(r) and r[i] is not None:
                    # 여러 줄인 값은 가장 긴 한 줄 기준으로 계산 (셀 안에서 줄바꿈되므로)
                    longest = max(longest, max((len(line) for line in str(r[i]).split("\n")), default=0))
            widths.append(min(max(longest * per_char + pad, min_w), max_w))
        return widths

    def shrink_font(table, size=Pt(8)):
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.font.size = size

    def add_table(document, headers, rows, max_rows=None):
        shown = rows if max_rows is None else rows[:max_rows]
        t = document.add_table(rows=1, cols=len(headers))
        t.style = "Light Grid Accent 1"
        t.alignment = WD_TABLE_ALIGNMENT.LEFT
        for cell, h in zip(t.rows[0].cells, headers):
            cell.text = str(h)
            cell.paragraphs[0].runs[0].bold = True
        for r in shown:
            cells = t.add_row().cells
            for cell, v in zip(cells, r):
                cell.text = "" if v is None else str(v)
        widths = compute_col_widths(headers, shown)
        set_table_fixed_layout(t)
        set_col_widths(t, widths)
        shrink_font(t)
        if max_rows is not None and len(rows) > max_rows:
            document.add_paragraph(f"... 외 {len(rows) - max_rows}행 더 있음 (전체 내용은 .xlsx로 출력해주세요)")
        return t

    document = docx.Document()
    for section in document.sections:
        section.page_width, section.page_height = section.page_height, section.page_width  # 가로 방향
        section.left_margin = section.right_margin = Pt(24)

    document.add_heading("스마트 자료취합기 - 통합결과", level=1)
    document.add_paragraph(f"총 {len(all_rows)}행, 표준 열 {len(canonical_columns)}개")
    if all_rows:
        headers = canonical_columns + ["출처파일"]
        rows = [[row_dict.get(c, "") for c in canonical_columns] + [source] for row_dict, source in all_rows]
        add_table(document, headers, rows, max_rows=200)
    else:
        document.add_paragraph("합쳐진 데이터가 없습니다.")

    flagged = [(label, orig, canon, method, score) for label, m in file_mappings
               for orig, (canon, method, score) in m.items() if method != "사전매칭"]
    if flagged:
        document.add_heading("열이름맞춤표 (확인이 필요할 수 있는 열만)", level=1)
        add_table(document, MAPPING_HEADER,
                  [[l, o, c, _method_label(m), s] for l, o, c, m, s in flagged], max_rows=200)

    if skipped:
        document.add_heading("처리못한 파일", level=1)
        for fname, reason in skipped:
            document.add_paragraph(f"{fname} — {reason}", style="List Bullet")

    document.save(output_path)


def write_pdf(output_path, canonical_columns, all_rows, file_mappings, skipped):
    """
    reportlab으로 직접 PDF를 만든다 (LibreOffice 같은 외부 프로그램이 필요 없음 --
    사용자 컴퓨터에 파이썬만 있으면 바로 동작하도록 하기 위함).
    한글 표시를 위해 reportlab에 내장된 CID 폰트(HYGothic-Medium)를 사용한다.
    """
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4, landscape
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    # 한글 폰트를 두 단계로 시도한다.
    # 1) 윈도우에 기본으로 깔려있는 맑은 고딕(malgun.ttf) -- 실제 글자 모양을
    #    PDF 안에 통째로 넣어버리기 때문에(임베딩), 받는 사람 컴퓨터에 어떤
    #    프로그램으로 열어도 한글이 깨지지 않는다. 이성빈님 컴퓨터에서 확인.
    # 2) 위 폰트를 못 찾으면(윈도우가 아니거나 폰트가 없는 경우), reportlab에
    #    내장된 한글 폰트 이름표(HYGothic-Medium)를 대신 사용한다. 이 경우
    #    PDF를 여는 프로그램이 자체적으로 한글 글꼴을 갖고 있어야 정상적으로
    #    보인다(대부분의 최신 PDF 뷰어는 문제없지만, 100% 보장은 아니다).
    FONT = "HangulPDFFont"
    WINDOWS_FONT_CANDIDATES = [
        r"C:\Windows\Fonts\malgun.ttf",
        r"C:\Windows\Fonts\malgunbd.ttf",
        r"C:\Windows\Fonts\gulim.ttc",
    ]
    embedded = False
    for font_path in WINDOWS_FONT_CANDIDATES:
        if os.path.exists(font_path):
            try:
                pdfmetrics.registerFont(TTFont(FONT, font_path))
                embedded = True
                break
            except Exception:
                continue
    if not embedded:
        FONT = "HYGothic-Medium"
        pdfmetrics.registerFont(UnicodeCIDFont(FONT))
        _note("참고: 윈도우 기본 한글 폰트(맑은 고딕)를 찾지 못해 대체 폰트로 PDF를 만들었습니다.")
        _note("      PDF를 열어서 한글이 정상적으로 보이는지 한 번 확인해주세요.")
    body_style = ParagraphStyle("body", fontName=FONT, fontSize=8, leading=10)
    head_style = ParagraphStyle("head", fontName=FONT, fontSize=9, leading=11, textColor=colors.white)
    title_style = ParagraphStyle("title", fontName=FONT, fontSize=16, leading=20, spaceAfter=10)
    h2_style = ParagraphStyle("h2", fontName=FONT, fontSize=13, leading=16, spaceBefore=16, spaceAfter=8)

    def para(v):
        return Paragraph("" if v is None else str(v).replace("\n", "<br/>"), body_style)

    def make_table(headers, rows, max_rows=None):
        shown = rows if max_rows is None else rows[:max_rows]
        data = [[Paragraph(str(h), head_style) for h in headers]] + [[para(v) for v in r] for r in shown]
        t = Table(data, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1F4E78")),
            ("GRID", (0, 0), (-1, -1), 0.5, colors.grey),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#F2F2F2")]),
        ]))
        elems = [t]
        if max_rows is not None and len(rows) > max_rows:
            elems.append(Paragraph(f"... 외 {len(rows) - max_rows}행 더 있음 (전체 내용은 .xlsx로 출력해주세요)", body_style))
        return elems

    doc = SimpleDocTemplate(output_path, pagesize=landscape(A4), leftMargin=24, rightMargin=24, topMargin=24, bottomMargin=24)
    story = [Paragraph("스마트 자료취합기 - 통합결과", title_style),
             Paragraph(f"총 {len(all_rows)}행, 표준 열 {len(canonical_columns)}개", body_style)]

    if all_rows:
        headers = canonical_columns + ["출처파일"]
        rows = [[row_dict.get(c, "") for c in canonical_columns] + [source] for row_dict, source in all_rows]
        story.append(Spacer(1, 8))
        story += make_table(headers, rows, max_rows=150)
    else:
        story.append(Paragraph("합쳐진 데이터가 없습니다.", body_style))

    flagged = [(label, orig, canon, method, score) for label, m in file_mappings
               for orig, (canon, method, score) in m.items() if method != "사전매칭"]
    if flagged:
        story.append(Paragraph("열이름맞춤표 (확인이 필요할 수 있는 열만)", h2_style))
        story += make_table(MAPPING_HEADER,
                             [[l, o, c, _method_label(m), s] for l, o, c, m, s in flagged], max_rows=150)

    if skipped:
        story.append(Paragraph("처리못한 파일", h2_style))
        for fname, reason in skipped:
            story.append(Paragraph(f"• {fname} — {reason}", body_style))

    doc.build(story)


# ---------------------------------------------------------------------------
# 7. 한글(HWPX) 출력 -- 실제 hwpx 파일에서 "내용은 비우고 스타일/글꼴 정의만 남긴"
#    빈 틀(hwpx_template.hwpx, 스크립트와 같은 폴더에 들어있음)을 재사용해서,
#    그 안에 우리가 합친 표를 진짜 HWPX 표 구조(hp:tbl)로 새로 채워 넣는 방식이다.
#    HWPX는 스타일/글꼴 정의가 얽혀있는 구조라, 완전히 빈 문서에서부터 전부
#    새로 만드는 것보다 이미 유효하게 동작하는 파일의 틀을 재사용하는 편이
#    훨씬 안전하다 (실제 한글 프로그램에서 열어서 확인 완료).
# ---------------------------------------------------------------------------
if getattr(sys, "frozen", False):
    # PyInstaller로 exe 하나로 묶었을 때: --add-data로 함께 넣은 파일은
    # 실행 중 sys._MEIPASS(임시 압축해제 폴더) 안에 들어있다.
    _BASE_DIR = getattr(sys, "_MEIPASS", os.path.dirname(sys.executable))
else:
    _BASE_DIR = os.path.dirname(os.path.abspath(__file__))
_HWPX_TEMPLATE = os.path.join(_BASE_DIR, "hwpx_template.hwpx")


def _hwpx_esc(s):
    s = "" if s is None else str(s)
    return (s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
             .replace("\n", " / "))


def _hwpx_paragraph(text, para_pr="0", style="22", char_pr="0"):
    return (f'<hp:p id="0" paraPrIDRef="{para_pr}" styleIDRef="{style}" pageBreak="0" '
            f'columnBreak="0" merged="0"><hp:run charPrIDRef="{char_pr}">'
            f'<hp:t>{_hwpx_esc(text)}</hp:t></hp:run>'
            f'<hp:linesegarray><hp:lineseg textpos="0" vertpos="0" vertsize="1000" '
            f'textheight="1000" baseline="850" spacing="600" horzpos="0" horzsize="49500" '
            f'flags="393216"/></hp:linesegarray></hp:p>')


def _hwpx_char_pt(char_pr):
    """템플릿(hwpx_template.hwpx)의 charPr id별 실제 글자 크기(pt). 폭 계산에 사용."""
    return {"9": 10.0, "2": 9.0, "0": 10.0}.get(char_pr, 10.0)


def _hwpx_text_width(text, font_pt):
    """텍스트가 실제로 차지하는 대략적인 가로 폭(HWPUNIT, 1pt = 100 HWPUNIT).
    한글/한자 등 전각 문자는 글자크기만큼, 영문/숫자/공백 등 반각 문자는 절반만큼 잡는다."""
    width = 0.0
    for ch in text:
        code = ord(ch)
        is_wide = (0xAC00 <= code <= 0xD7A3) or (0x4E00 <= code <= 0x9FFF) or (0x3130 <= code <= 0x318F)
        width += font_pt * 100 * (1.0 if is_wide else 0.55)
    return width


def _hwpx_wrap_lines(text, col_width, font_pt):
    """실제 한글 프로그램이 셀 안에서 자동으로 줄바꿈해줄 거라고 믿지 않고,
    우리가 직접 글자 폭을 계산해서 줄을 미리 나눈다 — 문단(<hp:p>)을 줄 개수만큼
    만들어서, 한글이 레이아웃을 다시 계산하지 않고 우리가 적어둔 값을 그대로 써도
    안전하게 보이도록 하기 위함이다 (실제 파일로 검증한 결과, 한글은 셀 폭에 맞춰
    자동으로 다시 줄바꿈해주지 않고 우리가 적어둔 값을 그대로 신뢰하는 것으로 확인됨).
    빈 문자열이면 빈 줄 1개를 반환한다."""
    usable = max(col_width - 400, 300)
    lines = []
    # 주의: text가 0 / 0.0 처럼 파이썬에서 "거짓"으로 취급되는 값일 수 있으므로
    # `if text`가 아니라 None인지로만 판단한다 (0.0을 빈 칸으로 지워버리는 버그 방지).
    text_str = "" if text is None else str(text)
    forced_parts = text_str.split(" / ") if text_str != "" else [""]
    for part in forced_parts:
        if part == "":
            lines.append("")
            continue
        cur = ""
        cur_w = 0.0
        for ch in part:
            code = ord(ch)
            is_wide = (0xAC00 <= code <= 0xD7A3) or (0x4E00 <= code <= 0x9FFF) or (0x3130 <= code <= 0x318F)
            ch_w = font_pt * 100 * (1.0 if is_wide else 0.55)
            if cur and cur_w + ch_w > usable:
                lines.append(cur)
                cur, cur_w = ch, ch_w
            else:
                cur += ch
                cur_w += ch_w
        lines.append(cur)
    return lines or [""]


def _hwpx_lines_needed(text, col_width, font_pt):
    """주어진 칸 너비(HWPUNIT) 안에 이 글자를 넣으려면 몇 줄이 필요한지 계산."""
    return len(_hwpx_wrap_lines(text, col_width, font_pt))


def _hwpx_cell(text, col_addr, row_addr, width, height, char_pr="0"):
    font_pt = _hwpx_char_pt(char_pr)
    lines = _hwpx_wrap_lines(text, width, font_pt)
    paras = []
    for line in lines:
        line_w = max(int(_hwpx_text_width(line, font_pt)), 1)
        paras.append(
            f'<hp:p id="0" paraPrIDRef="20" styleIDRef="22" pageBreak="0" columnBreak="0" '
            f'merged="0"><hp:run charPrIDRef="{char_pr}"><hp:t>{_hwpx_esc(line)}</hp:t></hp:run>'
            f'<hp:linesegarray><hp:lineseg textpos="0" vertpos="0" vertsize="1000" '
            f'textheight="1000" baseline="850" spacing="600" horzpos="0" horzsize="{line_w}" '
            f'flags="393216"/></hp:linesegarray></hp:p>'
        )
    inner_p = "".join(paras)
    return (f'<hp:tc name="" header="0" hasMargin="0" protect="0" editable="0" dirty="0" '
            f'borderFillIDRef="4">'
            f'<hp:subList id="" textDirection="HORIZONTAL" lineWrap="BREAK" vertAlign="CENTER" '
            f'linkListIDRef="0" linkListNextIDRef="0" textWidth="0" textHeight="0" hasTextRef="0" '
            f'hasNumRef="0">{inner_p}</hp:subList>'
            f'<hp:cellAddr colAddr="{col_addr}" rowAddr="{row_addr}"/>'
            f'<hp:cellSpan colSpan="1" rowSpan="1"/>'
            f'<hp:cellSz width="{width}" height="{height}"/>'
            f'<hp:cellMargin left="200" right="200" top="100" bottom="100"/></hp:tc>')


def _hwpx_compute_col_widths(headers, rows, total_width, sample_limit=40):
    """열마다 내용 길이에 비례해서 너비를 나눠준다 (전체 합은 total_width로 고정).
    표가 좁아서 잘리거나(=열이 페이지 밖으로 넘침) 글자가 겹쳐 보이는 문제를 막기 위함."""
    col_cnt = len(headers)
    if col_cnt == 0:
        return []
    font_pt = _hwpx_char_pt("9")
    weights = []
    for c, h in enumerate(headers):
        w = _hwpx_text_width(str(h), font_pt)
        for row in rows[:sample_limit]:
            if c < len(row):
                w = max(w, _hwpx_text_width(str(row[c]), font_pt))
        weights.append(max(w, font_pt * 100))  # 최소 한 글자 폭은 확보

    min_w = 1800   # 어떤 열도 이보다 좁아지지는 않게 (약 0.63cm)
    max_w = int(total_width * 0.35)  # 한 열이 표 전체를 독차지하지 않게

    if min_w * col_cnt >= total_width:
        # 열이 너무 많아서 최소 너비조차 다 못 지킬 때는 그냥 균등하게 나눈다
        base = total_width // col_cnt
        col_widths = [base] * col_cnt
        col_widths[-1] += total_width - base * col_cnt
        return col_widths

    # "물 채우기" 방식: 최소/최대 폭에 걸리는 열은 그 값으로 고정하고, 나머지 열끼리
    # 남은 폭을 내용 길이 비율대로 다시 나누는 과정을 값이 안정될 때까지 반복한다.
    # (이렇게 해야 열이 아무리 많아도 폭의 합이 total_width를 절대 넘지 않는다)
    col_widths = [0] * col_cnt
    fixed = [False] * col_cnt
    for _ in range(col_cnt + 1):
        free_idx = [i for i in range(col_cnt) if not fixed[i]]
        if not free_idx:
            break
        used = sum(col_widths[i] for i in range(col_cnt) if fixed[i])
        remaining = total_width - used
        free_weight = sum(weights[i] for i in free_idx) or 1
        newly_fixed = False
        for i in free_idx:
            share = remaining * weights[i] / free_weight
            if share < min_w:
                col_widths[i] = min_w
                fixed[i] = True
                newly_fixed = True
            elif share > max_w:
                col_widths[i] = max_w
                fixed[i] = True
                newly_fixed = True
        if not newly_fixed:
            for i in free_idx:
                col_widths[i] = int(remaining * weights[i] / free_weight)
            break

    # 반올림 오차를 가장 넓은 열에서 보정해 합이 정확히 total_width가 되도록 맞춘다
    diff = total_width - sum(col_widths)
    if diff:
        idx = max(range(col_cnt), key=lambda i: col_widths[i])
        col_widths[idx] += diff
    return col_widths


def _hwpx_table(headers, rows, table_id=1500000000):
    col_cnt = len(headers)
    row_cnt = 1 + len(rows)
    total_width = 49500
    col_widths = _hwpx_compute_col_widths(headers, rows, total_width)
    table_width = sum(col_widths)

    header_pt, data_pt = _hwpx_char_pt("9"), _hwpx_char_pt("2")
    line_h = {"9": 1350, "2": 1220}  # 줄 하나의 높이(글자크기 10pt/9pt 기준, 여유 포함)
    pad = 260  # 위아래 여백

    def row_height(row_texts, char_pr, font_pt):
        lines = max(_hwpx_lines_needed(t, col_widths[c], font_pt) for c, t in enumerate(row_texts)) if row_texts else 1
        return lines * line_h[char_pr] + pad

    trs = []
    header_h = row_height(headers, "9", header_pt)
    header_cells = "".join(_hwpx_cell(h, c, 0, col_widths[c], header_h, char_pr="9") for c, h in enumerate(headers))
    trs.append(f"<hp:tr>{header_cells}</hp:tr>")
    row_heights = [header_h]
    for r_idx, row in enumerate(rows, start=1):
        r_h = row_height(row, "2", data_pt)
        cells = "".join(_hwpx_cell(v, c, r_idx, col_widths[c], r_h, char_pr="2") for c, v in enumerate(row))
        trs.append(f"<hp:tr>{cells}</hp:tr>")
        row_heights.append(r_h)

    table_height = sum(row_heights)

    tbl = (f'<hp:tbl id="{table_id}" zOrder="1" numberingType="TABLE" textWrap="TOP_AND_BOTTOM" '
           f'textFlow="BOTH_SIDES" lock="0" dropcapstyle="None" pageBreak="NONE" repeatHeader="1" '
           f'rowCnt="{row_cnt}" colCnt="{col_cnt}" cellSpacing="0" borderFillIDRef="4" noAdjust="0">'
           f'<hp:sz width="{table_width}" widthRelTo="ABSOLUTE" height="{table_height}" '
           f'heightRelTo="ABSOLUTE" protect="0"/>'
           f'<hp:pos treatAsChar="1" affectLSpacing="0" flowWithText="1" allowOverlap="0" '
           f'holdAnchorAndSO="0" vertRelTo="PARA" horzRelTo="PARA" vertAlign="TOP" horzAlign="LEFT" '
           f'vertOffset="0" horzOffset="0"/>'
           f'<hp:outMargin left="141" right="141" top="141" bottom="141"/>'
           f'<hp:inMargin left="0" right="0" top="0" bottom="0"/>'
           + "".join(trs) + '</hp:tbl>')

    return (f'<hp:p id="0" paraPrIDRef="0" styleIDRef="22" pageBreak="0" columnBreak="0" merged="0">'
            f'<hp:run charPrIDRef="0">{tbl}</hp:run>'
            f'<hp:linesegarray><hp:lineseg textpos="0" vertpos="0" vertsize="1000" '
            f'textheight="1000" baseline="850" spacing="600" horzpos="0" horzsize="49500" '
            f'flags="393216"/></hp:linesegarray></hp:p>')


def write_hwpx(output_path, canonical_columns, all_rows, file_mappings, skipped):
    if not os.path.exists(_HWPX_TEMPLATE):
        _note(f"오류: hwpx_template.hwpx 파일을 찾을 수 없습니다. smart_merge_v2.py와 같은 폴더에 있어야 합니다.")
        _note(f"      (찾은 위치: {_HWPX_TEMPLATE})")
        sys.exit(1)

    with zipfile.ZipFile(_HWPX_TEMPLATE) as zin:
        template_section = zin.read("Contents/section0.xml").decode("utf-8")
        sec_start = template_section.find("<hs:sec")
        sec_open_end = template_section.find(">", sec_start) + 1
        sec_open_tag = template_section[sec_start:sec_open_end]
        first_p_start = template_section.find("<hp:p ")
        first_p_end = template_section.find("</hp:p>", first_p_start) + len("</hp:p>")
        first_p_block = template_section[first_p_start:first_p_end]

        parts = [sec_open_tag, first_p_block]
        parts.append(_hwpx_paragraph("스마트 자료취합기 - 통합결과"))
        parts.append(_hwpx_paragraph(f"총 {len(all_rows)}행, 표준 열 {len(canonical_columns)}개"))

        MAX_ROWS = 300  # HWPX 표 하나가 너무 커지면 한글에서 열 때 느려질 수 있어 안전하게 제한
        if all_rows:
            headers = canonical_columns + ["출처파일"]
            rows = [[row_dict.get(c, "") for c in canonical_columns] + [source] for row_dict, source in all_rows]
            parts.append(_hwpx_table(headers, rows[:MAX_ROWS]))
            if len(rows) > MAX_ROWS:
                parts.append(_hwpx_paragraph(f"... 외 {len(rows) - MAX_ROWS}행 더 있음 (전체 내용은 .xlsx로 출력해주세요)"))
        else:
            parts.append(_hwpx_paragraph("합쳐진 데이터가 없습니다."))

        flagged = [(label, orig, canon, method, score) for label, m in file_mappings
                   for orig, (canon, method, score) in m.items() if method != "사전매칭"]
        if flagged:
            parts.append(_hwpx_paragraph("열이름맞춤표 (확인이 필요할 수 있는 열만)"))
            parts.append(_hwpx_table(
                MAPPING_HEADER,
                [[l, o, c, _method_label(m), s] for l, o, c, m, s in flagged][:MAX_ROWS],
                table_id=1500000001,
            ))

        if skipped:
            parts.append(_hwpx_paragraph("처리못한 파일"))
            for fname, reason in skipped:
                parts.append(_hwpx_paragraph(f"- {fname} : {reason}"))

        new_section = '<?xml version="1.0" encoding="UTF-8" standalone="yes" ?>' + "".join(parts) + "</hs:sec>"

        with zipfile.ZipFile(output_path, "w") as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename == "Contents/section0.xml":
                    data = new_section.encode("utf-8")
                compress_type = zipfile.ZIP_STORED if item.compress_type == 0 else zipfile.ZIP_DEFLATED
                zi = zipfile.ZipInfo(item.filename, date_time=item.date_time)
                zi.compress_type = compress_type
                zout.writestr(zi, data)


WRITERS = {
    ".xlsx": write_xlsx,
    ".csv": write_csv,
    ".docx": write_docx,
    ".pdf": write_pdf,
    ".hwpx": write_hwpx,
}


# ---------------------------------------------------------------------------
# 6. 진입점 -- 출력 파일의 확장자를 보고 형식을 자동으로 고른다.
# ---------------------------------------------------------------------------
def smart_merge(input_dir, output_path):
    ext = os.path.splitext(output_path)[1].lower()
    if ext not in WRITERS:
        supported = ", ".join(WRITERS.keys())
        _note(f"지원하지 않는 출력 형식입니다: '{ext}'. 다음 중 하나로 저장해주세요: {supported}")
        sys.exit(1)

    canonical_columns, all_rows, file_mappings, skipped = merge_all(input_dir)
    WRITERS[ext](output_path, canonical_columns, all_rows, file_mappings, skipped)

    _note(f"완료: {len(file_mappings)}개 표 -> {len(all_rows)}행 -> {output_path}")
    _note(f"표준 열 {len(canonical_columns)}개: {canonical_columns}")
    flagged = [(f, o) for f, m in file_mappings for o, (c, method, s) in m.items() if method == "신규열(수동확인필요)"]
    if flagged:
        _note(f"⚠ 수동 확인이 필요한 열 {len(flagged)}개")
    if skipped:
        _note(f"⚠ 처리하지 못한 파일 {len(skipped)}개")
        for fname, reason in skipped:
            _note(f"   - {fname}: {reason}")


if __name__ == "__main__":
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
    if len(sys.argv) != 3:
        supported = ", ".join(WRITERS.keys())
        _note(f"사용법: python smart_merge_v2.py <입력폴더> <출력파일.xlsx|.docx|.csv|.pdf|.hwpx>")
        _note(f"지원하는 출력 형식: {supported}")
        sys.exit(1)
    smart_merge(sys.argv[1], sys.argv[2])
